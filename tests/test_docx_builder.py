"""Tests for :mod:`clinic.printing.docx_builder`."""

from __future__ import annotations

from datetime import datetime
from pathlib import Path

import pytest
from docx import Document

from clinic.domain import doctor_service, reception_service
from clinic.domain.dto import PatientInput, ReceptionInput
from clinic.printing.docx_builder import (
    build_reception_document,
    save_reception_document,
)


@pytest.fixture
def sample() -> dict:
    doc = doctor_service.create(full_name="Karimov Ali", phone="+998901234567")
    reception, patient, _ = reception_service.save(
        ReceptionInput(
            patient=PatientInput(
                full_name="Aliyev Anvar",
                birth_year=1990,
                phone="+998901112233",
                address="Toshkent, Yunusobod",
            ),
            patient_id=None,
            doctor_id=doc.id,
            reception_date=datetime(2026, 7, 12, 10, 30),
            complaints_codes=["ear_pain", "ear_discharge"],
            complaints_details={"ear_discharge": "purulent"},
            complaints_note="3 kundan beri",
            anamnesis="5 kun oldin boshlangan.",
            lor_status={
                "rhinoscopy": {"breathing": {"state": "free"}},
                "otoscopy": {
                    "AS": {
                        "tympanic_membrane": {
                            "color": "hyperemic",
                            "perforation": "central",
                        }
                    }
                },
            },
            diagnosis="Otitis media akuta",
            recommendation="Antibiotik + kompress",
        )
    )
    return {
        "doctor": doctor_service.get(doc.id),
        "reception": reception,
        "patient": patient,
        "clinic": {
            "name_uz": "LOR klinikasi",
            "name_ru": "ЛОР клиника",
            "address_uz": "Toshkent, Yunusobod",
            "address_ru": "Ташкент, Юнусабад",
            "phone": "+998 71 XX XX XX",
        },
    }


def _all_text(doc) -> str:  # type: ignore[no-untyped-def]
    """Flatten all body text — including cells of any tables — into one blob."""
    parts = [p.text for p in doc.paragraphs]
    for table in getattr(doc, "tables", []):
        for row in table.rows:
            for cell in row.cells:
                parts.extend(p.text for p in cell.paragraphs)
    return "\n".join(parts)


def test_reception_default_uz(sample: dict, tmp_path: Path) -> None:
    # Force the default renderer (bypass any shipped template).
    doc = build_reception_document(
        reception=sample["reception"],
        patient=sample["patient"],
        doctor=sample["doctor"],
        clinic=sample["clinic"],
        lang="uz",
        template_path=tmp_path / "nonexistent.docx",
    )
    text = _all_text(doc)
    assert "LOR klinikasi" in text
    # Phase 9: Uzbek heading is now Cyrillic per user preference.
    assert "ҚАБУЛ ВАРАҚАСИ" in text
    assert "Aliyev Anvar" in text
    assert "Otitis media akuta" in text
    assert "қулоқда оғриқ" in text.lower()
    assert "LOR STATUS" in text  # stays in Latin


def test_reception_default_ru(sample: dict, tmp_path: Path) -> None:
    doc = build_reception_document(
        reception=sample["reception"],
        patient=sample["patient"],
        doctor=sample["doctor"],
        clinic=sample["clinic"],
        lang="ru",
        template_path=tmp_path / "nonexistent.docx",
    )
    text = _all_text(doc)
    assert "ЛОР клиника" in text
    assert "ЛИСТ ПРИЁМА" in text
    # LOR STATUS stays Latin in both languages per Phase 9 request.
    assert "LOR STATUS" in text


def test_reception_omits_optional_sections(sample: dict, tmp_path: Path) -> None:
    """Recommendation/anamnesis are optional — no heading when empty."""
    reception = sample["reception"]
    reception.recommendation = None
    reception.anamnesis = None

    doc = build_reception_document(
        reception=reception,
        patient=sample["patient"],
        doctor=sample["doctor"],
        clinic=sample["clinic"],
        lang="uz",
        template_path=tmp_path / "nonexistent.docx",
    )
    text = _all_text(doc)
    # No RECOMMENDATION / ANAMNEZ heading should appear.
    assert "ТАВСИЯ" not in text
    assert "АНАМНЕЗ" not in text


def test_save_reception_document_writes_valid_file(
    sample: dict, tmp_path: Path
) -> None:
    dest = tmp_path / "out.docx"
    # Explicitly bypass the shipped ``reception_template.docx`` so we're
    # exercising the default renderer's output on disk.
    save_reception_document(
        output_path=dest,
        reception=sample["reception"],
        patient=sample["patient"],
        doctor=sample["doctor"],
        clinic=sample["clinic"],
        lang="uz",
        template_path=tmp_path / "nonexistent.docx",
    )
    assert dest.exists()
    # Round-trip: re-open the file.
    doc = Document(str(dest))
    assert any("ҚАБУЛ ВАРАҚАСИ" in p.text for p in doc.paragraphs)


def test_reception_with_user_template(sample: dict, tmp_path: Path) -> None:
    """When a template with Jinja placeholders is provided, docxtpl fills it in."""
    # Build a minimal template on the fly.
    template_path = tmp_path / "reception_template.docx"
    tpl = Document()
    tpl.add_paragraph("Custom title: {{ clinic.name }}")
    tpl.add_paragraph("Bemor: {{ patient.full_name }} ({{ patient.age }})")
    tpl.add_paragraph("Tashxis: {{ reception.diagnosis }}")
    tpl.save(str(template_path))

    doc = build_reception_document(
        reception=sample["reception"],
        patient=sample["patient"],
        doctor=sample["doctor"],
        clinic=sample["clinic"],
        lang="uz",
        template_path=template_path,
    )
    text = _all_text(doc)
    # The placeholders should have been substituted, not left literal.
    assert "{{" not in text
    assert "LOR klinikasi" in text
    assert "Aliyev Anvar" in text
    assert "Otitis media akuta" in text


def test_missing_template_falls_back_to_default(sample: dict, tmp_path: Path) -> None:
    ghost = tmp_path / "not_here.docx"
    doc = build_reception_document(
        reception=sample["reception"],
        patient=sample["patient"],
        doctor=sample["doctor"],
        clinic=sample["clinic"],
        lang="uz",
        template_path=ghost,
    )
    text = _all_text(doc)
    # Default header wins.
    assert "ҚАБУЛ ВАРАҚАСИ" in text
