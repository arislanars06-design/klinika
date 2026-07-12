"""Tests for :mod:`clinic.printing.stats_export`."""

from __future__ import annotations

from datetime import date, datetime, timedelta
from decimal import Decimal
from pathlib import Path

from docx import Document

from clinic.domain import (
    cashier_service,
    doctor_service,
    reception_service,
    service_service,
    stats_service,
)
from clinic.domain.dto import (
    CashierItemInput,
    CashierPaymentInput,
    PatientInput,
    ReceptionInput,
)
from clinic.domain.stats_service import PeriodPreset
from clinic.printing.stats_export import (
    build_cashier_stats_document,
    build_patient_stats_document,
    save_cashier_stats,
    save_patient_stats,
)


def _bootstrap() -> None:
    doc = doctor_service.create(full_name="Karimov Ali")
    svc = service_service.create(
        name_uz="Konsultatsiya", name_ru="Консультация", price=Decimal("100000")
    )
    reception, patient, _ = reception_service.save(
        ReceptionInput(
            patient=PatientInput(full_name="Aliyev Anvar", birth_year=1990),
            patient_id=None,
            doctor_id=doc.id,
            reception_date=datetime.now(),
            complaints_codes=["ear_pain"],
            diagnosis="Otitis media",
        )
    )
    cashier_service.save_payment(
        CashierPaymentInput(
            patient_id=patient.id,
            reception_id=reception.id,
            items=[CashierItemInput(service_id=svc.id, quantity=2)],
        )
    )


def _text(doc) -> str:  # type: ignore[no-untyped-def]
    return "\n".join(p.text for p in doc.paragraphs)


def test_patient_stats_default_layout() -> None:
    _bootstrap()
    period = stats_service.build_period(PeriodPreset.MONTH)
    stats = stats_service.patient_stats(period)
    doc = build_patient_stats_document(stats, period, clinic={"name_uz": "LOR"}, lang="uz")
    text = _text(doc)
    assert "KLINIKA STATISTIKASI" in text
    # Three tables: KPI row + TOP diagnoses + by-day
    assert len(doc.tables) == 3


def _all_text(doc) -> str:  # type: ignore[no-untyped-def]
    """Concat paragraphs + every table cell text (python-docx keeps them apart)."""
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    parts.append(p.text)
    return "\n".join(parts)


def test_cashier_stats_default_layout() -> None:
    _bootstrap()
    today = date.today()
    period = stats_service.build_custom(today - timedelta(days=1), today + timedelta(days=1))
    stats = stats_service.cashier_stats(period)
    doc = build_cashier_stats_document(stats, period, clinic={"name_uz": "LOR"}, lang="uz")
    text = _all_text(doc)
    assert "KASSA HISOBOTI" in text
    # KPI + by_service + by_day = 3 tables.
    assert len(doc.tables) == 3
    # Total revenue is 200 000 (100 000 * 2 quantity).
    assert "200 000" in text


def test_save_patient_stats_roundtrip(tmp_path: Path) -> None:
    _bootstrap()
    period = stats_service.build_period(PeriodPreset.MONTH)
    stats = stats_service.patient_stats(period)
    dest = tmp_path / "p_stats.docx"
    save_patient_stats(
        output_path=dest,
        stats=stats,
        period=period,
        clinic={"name_uz": "LOR"},
        lang="uz",
    )
    assert dest.exists()
    reopened = Document(str(dest))
    assert any("KLINIKA STATISTIKASI" in p.text for p in reopened.paragraphs)


def test_save_cashier_stats_roundtrip(tmp_path: Path) -> None:
    _bootstrap()
    today = date.today()
    period = stats_service.build_custom(today - timedelta(days=1), today + timedelta(days=1))
    stats = stats_service.cashier_stats(period)
    dest = tmp_path / "c_stats.docx"
    save_cashier_stats(
        output_path=dest,
        stats=stats,
        period=period,
        clinic={"name_uz": "LOR"},
        lang="uz",
    )
    assert dest.exists()
    reopened = Document(str(dest))
    assert any("KASSA HISOBOTI" in p.text for p in reopened.paragraphs)


def test_cashier_stats_empty_period_renders_placeholder() -> None:
    _bootstrap()
    period = stats_service.build_custom(date(2020, 1, 1), date(2020, 1, 31))
    stats = stats_service.cashier_stats(period)
    doc = build_cashier_stats_document(stats, period, clinic={"name_uz": "LOR"}, lang="uz")
    text = _text(doc)
    # "Ma'lumot yo'q" should appear where the empty tables would go.
    assert "yo'q" in text.lower() or "нет" in text.lower()
