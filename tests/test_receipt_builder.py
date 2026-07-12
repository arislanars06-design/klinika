"""Tests for :mod:`clinic.printing.receipt_builder`."""

from __future__ import annotations

from datetime import datetime
from decimal import Decimal
from pathlib import Path

import pytest
from docx import Document

from clinic.domain import (
    cashier_service,
    doctor_service,
    reception_service,
    service_service,
)
from clinic.domain.dto import (
    CashierItemInput,
    CashierPaymentInput,
    PatientInput,
    ReceptionInput,
)
from clinic.printing.receipt_builder import (
    build_receipt_document,
    save_receipt_document,
)


@pytest.fixture
def receipt_data() -> dict:
    doc = doctor_service.create(full_name="Karimov Ali")
    s1 = service_service.create(
        name_uz="Konsultatsiya", name_ru="Консультация", price=Decimal("100000")
    )
    s2 = service_service.create(
        name_uz="Audiometriya", name_ru="Аудиометрия", price=Decimal("150000")
    )
    reception, patient, _ = reception_service.save(
        ReceptionInput(
            patient=PatientInput(full_name="Aliyev Anvar", birth_year=1990),
            patient_id=None,
            doctor_id=doc.id,
            reception_date=datetime(2026, 7, 12, 10, 30),
            complaints_codes=["ear_pain"],
            diagnosis="Otitis",
        )
    )
    records = cashier_service.save_payment(
        CashierPaymentInput(
            patient_id=patient.id,
            reception_id=reception.id,
            items=[
                CashierItemInput(service_id=s1.id, quantity=1),
                CashierItemInput(service_id=s2.id, quantity=2),
            ],
            note="test note",
        )
    )
    return {
        "records": records,
        "patient": patient,
        "clinic": {"name_uz": "LOR", "phone": "+998 71 XX"},
    }


def _text(doc) -> str:  # type: ignore[no-untyped-def]
    return "\n".join(p.text for p in doc.paragraphs)


def test_receipt_default_layout(receipt_data: dict) -> None:
    doc = build_receipt_document(
        records=receipt_data["records"],
        patient=receipt_data["patient"],
        clinic=receipt_data["clinic"],
        lang="uz",
    )
    text = _text(doc)
    assert "KVITANSIYA" in text
    assert "Aliyev Anvar" in text
    assert "JAMI" in text
    # Grand total (100k + 300k = 400k)
    assert "400 000" in text
    # Table has 1 header row + 2 item rows.
    assert len(doc.tables) == 1
    assert len(doc.tables[0].rows) == 3


def test_receipt_ru_labels(receipt_data: dict) -> None:
    doc = build_receipt_document(
        records=receipt_data["records"],
        patient=receipt_data["patient"],
        clinic=receipt_data["clinic"],
        lang="ru",
    )
    text = _text(doc)
    assert "КВИТАНЦИЯ" in text
    assert "ИТОГО" in text


def test_receipt_save_roundtrip(receipt_data: dict, tmp_path: Path) -> None:
    dest = tmp_path / "receipt.docx"
    save_receipt_document(
        output_path=dest,
        records=receipt_data["records"],
        patient=receipt_data["patient"],
        clinic=receipt_data["clinic"],
        lang="uz",
    )
    assert dest.exists()
    reopened = Document(str(dest))
    assert any("KVITANSIYA" in p.text for p in reopened.paragraphs)


def test_receipt_requires_records(receipt_data: dict) -> None:
    with pytest.raises(ValueError):
        build_receipt_document(
            records=[],
            patient=receipt_data["patient"],
            clinic=receipt_data["clinic"],
            lang="uz",
        )
