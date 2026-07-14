"""Regression tests for the four immutable rules of the shipped reception
template.  These rules were requested by the operator and must survive
every future edit of ``scripts/build_reception_template.py``.

  1. Reception date appears exactly once (under the title).  No №<id>
     next to it, no second date in the footer.
  2. Rendered .docx has no visible content tables.  Only the borderless
     letterhead (1x3) and the boxed TAVSIYA (1x1) are allowed.
  3. There is no date next to the doctor's signature line.
  4. Only the TAVSIYA section is framed.

Every check works on the actual .docx produced by
``build_reception_document`` with realistic sample data — the same code
path that runs when the operator clicks *Save and Print* in production.
"""

from __future__ import annotations

import re
from datetime import datetime
from pathlib import Path

import pytest
from docx import Document

from clinic.domain.dto import DoctorDTO, PatientDTO, ReceptionDTO
from clinic.printing.docx_builder import build_reception_document

REPO_ROOT = Path(__file__).resolve().parents[1]
TEMPLATE = REPO_ROOT / "templates" / "reception_template.docx"


# ---------------------------------------------------------------------------
# Fixtures — a fully-populated reception rendered through the real pipeline
# ---------------------------------------------------------------------------

_CLINIC = {
    "name_uz": "NurMed LOR",
    "name_ru": "Клиника NurMed",
    "address_uz": "Toshkent, Sergeli",
    "address_ru": "Ташкент, Сергели",
    "phone": "+998 99 981 91 64",
    "logo_path": "",
}

FIXED_DATE = datetime(2026, 7, 14, 13, 19)


def _make_dtos():
    patient = PatientDTO(
        id=1,
        full_name="Musulmonov Rustam",
        birth_year=1998,
        address="Toshkent, Sergeli",
        phone="+998 90 555 44 33",
        created_at=FIXED_DATE,
        updated_at=FIXED_DATE,
    )
    doctor = DoctorDTO(
        id=1,
        full_name="Madirimov Xikmatulla",
        phone="+998 99 981 91 64",
        is_active=True,
        save_folder=None,
    )
    reception = ReceptionDTO(
        id=42,
        patient_id=patient.id,
        doctor_id=doctor.id,
        reception_date=FIXED_DATE,
        complaints_codes=[],
        complaints_details={},
        complaints_note="Bosh og'rig'i",
        anamnesis="3 kundan beri",
        lor_status={"rhinoscopy": "Tashqi burun: O'zgarmagan."},
        diagnosis="O'tkir tonzillofaringit",
        recommendation="Debara 1 tab * 4 mahal 5 kun.",
        created_at=FIXED_DATE,
    )
    return reception, patient, doctor


@pytest.fixture(scope="module")
def rendered_doc(tmp_path_factory) -> Document:
    """Render the shipped template with sample data and return the Document."""
    if not TEMPLATE.exists():
        pytest.skip(f"Shipped template not present at {TEMPLATE}")

    reception, patient, doctor = _make_dtos()
    doc = build_reception_document(
        reception=reception, patient=patient, doctor=doctor,
        clinic=_CLINIC, lang="uz",
    )
    out = tmp_path_factory.mktemp("rendered") / "reception.docx"
    doc.save(str(out))
    return Document(str(out))


def _all_text(doc: Document) -> str:
    """Flatten every paragraph (including inside tables) into one string."""
    parts: list[str] = []
    for para in doc.paragraphs:
        parts.append(para.text)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    parts.append(p.text)
    return "\n".join(parts)


# ---------------------------------------------------------------------------
# Rule 1 — one date, no №<id>
# ---------------------------------------------------------------------------

def test_rule1_reception_date_appears_exactly_once(rendered_doc):
    text = _all_text(rendered_doc)
    hits = re.findall(r"14\.07\.2026", text)
    assert len(hits) == 1, (
        f"Expected reception date to appear once, got {len(hits)}.\n"
        f"Full text:\n{text}"
    )


def test_rule1_no_reception_id_marker(rendered_doc):
    text = _all_text(rendered_doc)
    # The template no longer emits '№{{ reception.id }}' next to the date.
    # Fixture uses reception.id = 42, so any '№42' or standalone '№' after
    # the date means the marker leaked back in.
    assert "№42" not in text, "Reception №<id> marker leaked back into the sheet"
    # Also verify the raw placeholder isn't present (unrendered docxtpl).
    assert "{{ reception.id }}" not in text


def test_rule1_no_today_placeholder_in_template():
    """The template file itself must never emit {{ today }}."""
    import zipfile
    with zipfile.ZipFile(TEMPLATE) as z:
        body = z.read("word/document.xml").decode("utf-8", errors="replace")
    assert "{{ today }}" not in body
    assert "{{today}}" not in body


# ---------------------------------------------------------------------------
# Rule 2 — no visible content tables
# ---------------------------------------------------------------------------

def test_rule2_only_letterhead_and_tavsiya_tables(rendered_doc):
    """Every content section (patient info, signature footer) must be
    plain paragraphs.  The only tables allowed are:

      - a 1x3 letterhead (borderless — image + text + QR)
      - a 1x1 TAVSIYA frame (rule 4)
    """
    letterhead_seen = 0
    tavsiya_seen = 0
    for tbl in rendered_doc.tables:
        n_rows, n_cols = len(tbl.rows), len(tbl.columns)
        text = "\n".join(
            p.text for row in tbl.rows for cell in row.cells for p in cell.paragraphs
        )
        if n_rows == 1 and n_cols == 3:
            letterhead_seen += 1
        elif n_rows == 1 and n_cols == 1 and (
            "TAVSIYA" in text or "РЕКОМЕНДАЦИЯ" in text
        ):
            tavsiya_seen += 1
        else:
            pytest.fail(
                f"Unexpected {n_rows}x{n_cols} table found "
                f"(likely patient K/V or signature block):\n{text}"
            )

    assert letterhead_seen == 1
    assert tavsiya_seen == 1


# ---------------------------------------------------------------------------
# Rule 3 — no date next to the doctor's signature
# ---------------------------------------------------------------------------

def test_rule3_no_date_near_doctor_signature(rendered_doc):
    """The doctor signature line + 'Imzo' must not be followed by another
    date.  We look at every paragraph after the one containing the
    doctor's full name and confirm no reception-date-shaped string
    appears until end of document.
    """
    reception, patient, doctor = _make_dtos()
    doctor_name = doctor.full_name

    paragraphs = list(rendered_doc.paragraphs)
    hit_doctor = False
    for para in paragraphs:
        if doctor_name in para.text:
            hit_doctor = True
            continue
        if hit_doctor:
            # After the doctor's name, no date-shaped string is allowed.
            assert not re.search(r"\d{2}\.\d{2}\.\d{4}", para.text), (
                f"Date leaked into a paragraph after the doctor's name: "
                f"{para.text!r}"
            )


# ---------------------------------------------------------------------------
# Rule 4 — only TAVSIYA is framed
# ---------------------------------------------------------------------------

def test_rule4_only_tavsiya_has_visible_border(rendered_doc):
    """Walk every table cell in the rendered .docx and check its
    ``w:tcBorders`` XML.  A cell counts as 'visibly framed' if it has at
    least one non-'nil' border edge.  Exactly one framed cell is
    allowed — the TAVSIYA one.
    """
    from docx.oxml.ns import qn

    framed_cells = []
    for tbl in rendered_doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                tc_pr = cell._tc.tcPr
                if tc_pr is None:
                    continue
                borders = tc_pr.find(qn("w:tcBorders"))
                if borders is None:
                    continue
                has_visible = False
                for edge in ("top", "left", "bottom", "right"):
                    el = borders.find(qn(f"w:{edge}"))
                    if el is None:
                        continue
                    val = el.get(qn("w:val"))
                    if val and val != "nil" and val != "none":
                        has_visible = True
                        break
                if has_visible:
                    body = "\n".join(p.text for p in cell.paragraphs)
                    framed_cells.append(body)

    assert len(framed_cells) == 1, (
        f"Expected exactly one framed cell (TAVSIYA), got {len(framed_cells)}.\n"
        f"Framed cell contents:\n"
        + "\n---\n".join(framed_cells)
    )
    body = framed_cells[0]
    assert "TAVSIYA" in body or "РЕКОМЕНДАЦИЯ" in body, (
        f"The single framed cell is not the TAVSIYA one:\n{body}"
    )
