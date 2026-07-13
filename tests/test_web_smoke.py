"""Route-level smoke tests for the FastAPI web layer.

These use ``starlette.testclient.TestClient`` and drive the app just like a
real browser would: log in with a form POST, follow redirects, submit the
reception form, download the ``.docx``.

The db + templates directory are set up by ``conftest.py`` — nothing here
touches the real production data.
"""

from __future__ import annotations

import re

import pytest
from fastapi.testclient import TestClient

from clinic.domain import doctor_service
from clinic.web.app import create_app

PASSWORD = "clinic"  # default from WebSettings; overrideable via env


@pytest.fixture()
def client() -> TestClient:
    with TestClient(create_app()) as c:
        yield c


@pytest.fixture()
def logged_in(client: TestClient) -> TestClient:
    resp = client.post("/login", data={"password": PASSWORD}, follow_redirects=False)
    assert resp.status_code == 303
    return client


# ---------------------------------------------------------------------------
# Unauthenticated behaviour
# ---------------------------------------------------------------------------


def test_home_redirects_to_login_when_signed_out(client: TestClient) -> None:
    resp = client.get("/", follow_redirects=False)
    assert resp.status_code == 303
    assert "/login" in resp.headers["location"]


def test_login_page_renders(client: TestClient) -> None:
    resp = client.get("/login")
    assert resp.status_code == 200
    assert 'name="password"' in resp.text


def test_login_rejects_wrong_password(client: TestClient) -> None:
    resp = client.post("/login", data={"password": "wrong"}, follow_redirects=False)
    assert resp.status_code == 401


def test_login_accepts_correct_password(client: TestClient) -> None:
    resp = client.post("/login", data={"password": PASSWORD}, follow_redirects=False)
    assert resp.status_code == 303
    assert resp.headers["location"] == "/"


def test_login_next_redirect_is_same_origin(client: TestClient) -> None:
    resp = client.post(
        "/login",
        data={"password": PASSWORD, "next": "//evil.example.com/steal"},
        follow_redirects=False,
    )
    # Reject cross-origin ``next`` — must fall back to "/".
    assert resp.headers["location"] == "/"


# ---------------------------------------------------------------------------
# Authenticated pages
# ---------------------------------------------------------------------------


def test_home_renders_after_login(logged_in: TestClient) -> None:
    resp = logged_in.get("/")
    assert resp.status_code == 200
    assert "/reception/new" in resp.text
    assert "/patients" in resp.text


def test_patient_list_renders_empty(logged_in: TestClient) -> None:
    resp = logged_in.get("/patients")
    assert resp.status_code == 200


def test_reception_form_renders(logged_in: TestClient) -> None:
    resp = logged_in.get("/reception/new")
    assert resp.status_code == 200
    # Complaints accordion + LOR method tabs are present
    assert "complaints-accordion" in resp.text
    assert 'name="lor_rhinoscopy"' in resp.text
    assert 'name="doctor_id"' in resp.text


def test_language_switch_updates_html_lang(logged_in: TestClient) -> None:
    resp = logged_in.get("/lang/ru?next=/", follow_redirects=False)
    assert resp.status_code == 303
    resp = logged_in.get("/")
    assert 'lang="ru"' in resp.text


def test_logout_clears_session(logged_in: TestClient) -> None:
    resp = logged_in.get("/logout", follow_redirects=False)
    assert resp.status_code == 303
    resp = logged_in.get("/", follow_redirects=False)
    assert resp.status_code == 303  # back to login


# ---------------------------------------------------------------------------
# End-to-end: save reception → view → print
# ---------------------------------------------------------------------------


def test_full_reception_flow(logged_in: TestClient) -> None:
    doctor = doctor_service.create(
        full_name="Karimov Ali Valiyevich", phone="+998901234567"
    )

    form = {
        "full_name": "Aliyev Anvar Bahodirovich",
        "birth_year": "1990",
        "address": "Sergeli, 7-mavze",
        "phone": "+998939391914",
        "patient_id": "",
        "complaints": ["ear_pain", "ear_hearing_loss"],
        "complaints_note": "Kechqurun kuchayadi",
        "anamnesis": "Sovuq havodan keyin boshlangan.",
        "lor_rhinoscopy": "Shilliq parda giperemiyalangan.",
        "lor_pharyngoscopy": "",
        "lor_otoscopy": "O'ng quloq: pardasi giperemiyalangan.",
        "lor_laryngoscopy": "",
        "diagnosis": "O'tkir o'ng tomonlama otit (H66.9)",
        "recommendation": "Amoxicillin 500 mg 3 mahal 7 kun",
        "doctor_id": str(doctor.id),
    }
    resp = logged_in.post("/reception/new", data=form, follow_redirects=False)
    assert resp.status_code == 303, resp.text
    location = resp.headers["location"]
    assert location.startswith("/reception/")
    reception_id = int(location.rsplit("/", 1)[-1])

    # Detail page
    resp = logged_in.get(f"/reception/{reception_id}")
    assert resp.status_code == 200
    assert "Aliyev Anvar" in resp.text
    assert "otit" in resp.text.lower()
    assert "Karimov Ali" in resp.text
    assert "giperemiyalangan" in resp.text

    # Edit form pre-fills
    resp = logged_in.get(f"/reception/{reception_id}/edit")
    assert 'value="Aliyev Anvar Bahodirovich"' in resp.text
    assert 'value="1990"' in resp.text

    # Patient list surfaces the new patient
    resp = logged_in.get("/patients")
    assert resp.status_code == 200
    assert "Aliyev Anvar" in resp.text
    match = re.search(r'href="/patients/(\d+)"', resp.text)
    assert match, "no patient link in list HTML"
    patient_id = int(match.group(1))

    # Patient detail
    resp = logged_in.get(f"/patients/{patient_id}")
    assert resp.status_code == 200
    assert "otit" in resp.text.lower()

    # Autocomplete
    resp = logged_in.get("/patients/autocomplete?q=Ali")
    assert resp.status_code == 200
    assert "Aliyev" in resp.text

    # Docx download
    resp = logged_in.get(f"/print/reception/{reception_id}.docx")
    assert resp.status_code == 200
    assert resp.headers["content-type"].startswith(
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    assert len(resp.content) > 10_000

    # Inspect the docx bytes to confirm placeholders were rendered
    import io

    from docx import Document

    doc = Document(io.BytesIO(resp.content))
    body = "\n".join(p.text for p in doc.paragraphs) + "\n" + "\n".join(
        cell.text for tbl in doc.tables for row in tbl.rows for cell in row.cells
    )
    assert "Aliyev Anvar" in body
    assert "otit" in body.lower()
    assert "Karimov Ali" in body


def test_reception_form_validation_returns_400(logged_in: TestClient) -> None:
    """Empty required fields should render the form again with errors, not save."""
    resp = logged_in.post(
        "/reception/new",
        data={"full_name": "", "diagnosis": "", "complaints": []},
        follow_redirects=False,
    )
    assert resp.status_code == 400
    # Form is re-rendered — the accordion still appears
    assert "complaints-accordion" in resp.text
