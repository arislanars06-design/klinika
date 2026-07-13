"""Phase-2 web tests: cashier flow + stats dashboards + Word exports."""

from __future__ import annotations

import io

import pytest
from docx import Document
from fastapi.testclient import TestClient

from clinic.domain import doctor_service, service_service
from clinic.web.app import create_app

PASSWORD = "clinic"


@pytest.fixture()
def client() -> TestClient:
    with TestClient(create_app()) as c:
        c.post("/login", data={"password": PASSWORD}, follow_redirects=False)
        yield c


@pytest.fixture()
def seed(client: TestClient):
    """Create a doctor, two services, a patient (via reception form), and a
    persisted reception. Return the identifiers as a dict."""
    doctor = doctor_service.create(full_name="Karimov Ali", phone="+998901234567")
    s1 = service_service.create(name_uz="Konsultatsiya", name_ru="Консультация", price=100_000)
    s2 = service_service.create(name_uz="Audiometriya", name_ru="Аудиометрия", price=150_000)

    rec_form = {
        "full_name": "Aliyev Anvar",
        "birth_year": "1990",
        "address": "Sergeli",
        "phone": "+998939391914",
        "patient_id": "",
        "complaints": ["ear_pain"],
        "diagnosis": "Otit",
        "doctor_id": str(doctor.id),
    }
    resp = client.post("/reception/new", data=rec_form, follow_redirects=False)
    assert resp.status_code == 303
    reception_id = int(resp.headers["location"].rsplit("/", 1)[-1])

    # Find the auto-created patient id.
    resp = client.get(f"/reception/{reception_id}")
    import re
    m = re.search(r"/patients/(\d+)", resp.text)
    assert m
    patient_id = int(m.group(1))

    return {
        "doctor_id": doctor.id,
        "service_ids": [s1.id, s2.id],
        "reception_id": reception_id,
        "patient_id": patient_id,
    }


# ---------------------------------------------------------------------------
# Cashier landing + navigation
# ---------------------------------------------------------------------------


def test_cashier_landing_renders(client: TestClient) -> None:
    resp = client.get("/cashier")
    assert resp.status_code == 200
    assert "cashier" in resp.text.lower() or "kassa" in resp.text.lower() or "касс" in resp.text.lower()


def test_cashier_search_returns_matches(client: TestClient, seed: dict) -> None:
    resp = client.get("/cashier?q=Aliyev")
    assert resp.status_code == 200
    assert "Aliyev Anvar" in resp.text
    assert f"/cashier/patient/{seed['patient_id']}" in resp.text


def test_cashier_patient_page_shows_services_and_receptions(
    client: TestClient, seed: dict
) -> None:
    resp = client.get(f"/cashier/patient/{seed['patient_id']}")
    assert resp.status_code == 200
    # Cart JS + form scaffolding
    assert "services-json" in resp.text
    assert "cart-body" in resp.text
    # Reception is selectable
    assert f'value="{seed["reception_id"]}"' in resp.text


# ---------------------------------------------------------------------------
# Full save → receipt → print flow
# ---------------------------------------------------------------------------


def test_cashier_save_and_print_receipt(client: TestClient, seed: dict) -> None:
    save_form = {
        "service_id": [str(sid) for sid in seed["service_ids"]],
        "quantity": ["2", "1"],
        "reception_id": str(seed["reception_id"]),
        "note": "Konsultatsiya + audiometriya",
    }
    resp = client.post(
        f"/cashier/patient/{seed['patient_id']}/save",
        data=save_form,
        follow_redirects=False,
    )
    assert resp.status_code == 303
    loc = resp.headers["location"]
    assert loc.startswith("/cashier/receipt/")
    record_id = int(loc.rsplit("/", 1)[-1])

    # View receipt HTML
    resp = client.get(loc)
    assert resp.status_code == 200
    assert "Aliyev Anvar" in resp.text
    assert "Konsultatsiya" in resp.text or "Консультация" in resp.text
    assert "Audiometriya" in resp.text or "Аудиометрия" in resp.text

    # Download docx
    resp = client.get(f"/print/receipt/{record_id}.docx")
    assert resp.status_code == 200
    assert resp.headers["content-type"].startswith(
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    assert len(resp.content) > 5000

    doc = Document(io.BytesIO(resp.content))
    text = "\n".join(p.text for p in doc.paragraphs) + "\n" + "\n".join(
        c.text for tbl in doc.tables for row in tbl.rows for c in row.cells
    )
    assert "Aliyev Anvar" in text
    assert "KVITANSIYA" in text or "КВИТАНЦИЯ" in text
    assert "350" in text or "350000" in text  # grand total 100000*2 + 150000*1 = 350000


def test_cashier_empty_cart_is_rejected(client: TestClient, seed: dict) -> None:
    resp = client.post(
        f"/cashier/patient/{seed['patient_id']}/save",
        data={"reception_id": str(seed["reception_id"])},
        follow_redirects=False,
    )
    # Save handler flashes a validation-error and redirects back to the
    # patient page — no receipt is created.
    assert resp.status_code == 303
    assert resp.headers["location"] == f"/cashier/patient/{seed['patient_id']}"


# ---------------------------------------------------------------------------
# Stats dashboards
# ---------------------------------------------------------------------------


@pytest.mark.parametrize("path", [
    "/stats",
    "/stats?preset=today",
    "/stats?preset=week",
    "/stats?preset=month",
    "/stats?preset=year",
    "/stats?preset=custom&start=2026-01-01&end=2026-12-31",
    "/stats/cashier",
    "/stats/cashier?preset=today",
    "/stats/cashier?preset=custom&start=2026-01-01&end=2026-12-31",
])
def test_stats_pages_render(client: TestClient, path: str) -> None:
    resp = client.get(path)
    assert resp.status_code == 200
    # Common markers on all stats pages
    assert 'id="chart-data"' in resp.text


def test_stats_custom_period_rejects_bad_dates(client: TestClient) -> None:
    resp = client.get("/stats?preset=custom&start=not-a-date&end=2026-12-31")
    assert resp.status_code == 400


# ---------------------------------------------------------------------------
# Stats Word exports
# ---------------------------------------------------------------------------


def test_stats_patients_word_export(client: TestClient) -> None:
    resp = client.get("/print/stats/patients.docx?preset=month")
    assert resp.status_code == 200
    assert resp.headers["content-type"].startswith(
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    assert len(resp.content) > 5000
    doc = Document(io.BytesIO(resp.content))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "KLINIKA STATISTIKASI" in text or "СТАТИСТИКА" in text


def test_stats_cashier_word_export(client: TestClient) -> None:
    resp = client.get("/print/stats/cashier.docx?preset=month")
    assert resp.status_code == 200
    assert resp.headers["content-type"].startswith(
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    assert len(resp.content) > 5000
    doc = Document(io.BytesIO(resp.content))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "KASSA" in text.upper() or "КАСС" in text.upper()


# ---------------------------------------------------------------------------
# Nav integration — cashier is no longer disabled
# ---------------------------------------------------------------------------


def test_home_dashboard_has_active_cashier_card(client: TestClient) -> None:
    resp = client.get("/")
    assert resp.status_code == 200
    # Card is now a clickable link, not the opacity-50 placeholder.
    assert 'href="/cashier"' in resp.text
    # Stats dropdown appears in the nav
    assert "/stats/cashier" in resp.text
