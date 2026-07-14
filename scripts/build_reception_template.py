"""(Re)generate ``templates/reception_template.docx``.

A **compact** reception varaqasi:

  * A three-column header table:  [ logo ] [ clinic identity ] [ QR ]
    — both images are optional (missing files → cell stays empty).
  * Compact patient / complaints / anamnesis / LOR / diagnosis blocks
    with the section title inline (no big grey bars).
  * A **boxed** recommendation ("TAVSIYA") — a bordered table cell
    that stands out on the page, per the user's request.
  * Signature footer.

The template is a docxtpl document with Jinja placeholders; the clinic
identity fields come from ``ClinicInfo`` (Settings → Клиника).

Regenerate::

    python scripts/build_template_assets.py         # optional — refresh images
    python scripts/build_reception_template.py

The Settings → Шаблон page invokes ``build_default_template()`` directly
for the "Reset to default" button.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
DEFAULT_DST = ROOT / "templates" / "reception_template.docx"
ASSETS_DIR = ROOT / "templates" / "assets"

# Brand colours (fine to tweak; only affect text runs, not the images).
BRAND_BLUE_HEX = "0A3566"     # deep clinic-blue for divider + title
MUTED_GREY_HEX = "606060"     # meta text
BODY_BLACK = (0x11, 0x11, 0x11)


# ---------------------------------------------------------------------------
# Small XML helpers
# ---------------------------------------------------------------------------


def _set_cell_border(
    cell,
    *,
    size: int = 6,
    color: str = "808080",
    edges: tuple[str, ...] = ("top", "left", "bottom", "right"),
) -> None:
    """Apply a border to selected edges of a table cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.find(qn("w:tcBorders"))
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in edges:
        existing = tc_borders.find(qn(f"w:{edge}"))
        if existing is not None:
            tc_borders.remove(existing)
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(size))
        el.set(qn("w:color"), color)
        tc_borders.append(el)


def _no_cell_border(cell) -> None:
    """Explicitly disable all four edges (used inside the header layout)."""
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "nil")
        tc_borders.append(el)
    tc_pr.append(tc_borders)


def _shade_cell(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _add_horizontal_rule(doc, *, color: str = BRAND_BLUE_HEX, size: int = 12) -> None:
    """Draw a coloured horizontal rule under the current paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


# ---------------------------------------------------------------------------
# Reusable body components
# ---------------------------------------------------------------------------


def _inline_section(doc, title: str, placeholder: str, *, bold_body: bool = False,
                    body_size: int = 10) -> None:
    """One-line section: **TITLE:** placeholder.

    Keeps the sheet compact — no separate heading paragraph.
    """
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(1)
    run_t = p.add_run(f"{title}: ")
    run_t.bold = True
    run_t.font.size = Pt(10)
    run_t.font.color.rgb = RGBColor(*BODY_BLACK)
    run_b = p.add_run(placeholder)
    run_b.font.size = Pt(body_size)
    run_b.bold = bold_body


def _compact_patient_table(doc) -> None:
    """4-row two-column patient block; slimmer than the previous version."""
    tbl = doc.add_table(rows=2, cols=4)
    tbl.autofit = False
    widths = (Cm(3.0), Cm(6.5), Cm(3.0), Cm(4.5))
    for row in tbl.rows:
        for cell, width in zip(row.cells, widths, strict=True):
            cell.width = width
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            _set_cell_border(cell, size=4, color="B0B0B0")

    def _lbl(cell, text: str) -> None:
        _shade_cell(cell, "F0F4F8")
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

    def _val(cell, text: str) -> None:
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(text)
        r.font.size = Pt(10)

    _lbl(tbl.rows[0].cells[0], "F.I.SH.")
    _val(tbl.rows[0].cells[1], "{{ patient.full_name }}")
    _lbl(tbl.rows[0].cells[2], "Tug'ilgan y.")
    _val(tbl.rows[0].cells[3], "{{ patient.birth_year }} ({{ patient.age }} y)")

    _lbl(tbl.rows[1].cells[0], "Manzil")
    _val(tbl.rows[1].cells[1], "{{ patient.address }}")
    _lbl(tbl.rows[1].cells[2], "Telefon")
    _val(tbl.rows[1].cells[3], "{{ patient.phone }}")


def _boxed_recommendation(doc) -> None:
    """The one section that gets a visible framed box, per the user request.

    A one-cell table with a strong border around it and a light-blue tint
    inside.  Big enough to draw the eye but short enough to keep the sheet
    on one A4 page.
    """
    box = doc.add_table(rows=1, cols=1)
    box.autofit = False
    cell = box.rows[0].cells[0]
    cell.width = Cm(17.0)
    _set_cell_border(cell, size=14, color=BRAND_BLUE_HEX)
    _shade_cell(cell, "EAF3FA")

    # Title row inside the box
    p_title = cell.paragraphs[0]
    p_title.paragraph_format.space_before = Pt(2)
    p_title.paragraph_format.space_after = Pt(2)
    rt = p_title.add_run("TAVSIYA  /  РЕКОМЕНДАЦИЯ")
    rt.bold = True
    rt.font.size = Pt(11)
    rt.font.color.rgb = RGBColor(0x0A, 0x35, 0x66)

    # Body — placeholder
    p_body = cell.add_paragraph()
    p_body.paragraph_format.space_before = Pt(1)
    p_body.paragraph_format.space_after = Pt(3)
    rb = p_body.add_run("{{ reception.recommendation }}")
    rb.font.size = Pt(11)


# ---------------------------------------------------------------------------
# Header layout — [ logo ] [ clinic identity ] [ QR ]
# ---------------------------------------------------------------------------


def _write_letterhead(doc) -> None:
    """Three-cell header:  logo image | clinic text (name/addr/phone) | QR image.

    Images are pulled from ``templates/assets/{logo,qr}.png``.  Missing
    files are silently skipped — clinics without a logo still get a valid
    template.
    """
    hdr = doc.add_table(rows=1, cols=3)
    hdr.autofit = False
    logo_w, mid_w, qr_w = Cm(3.0), Cm(11.0), Cm(3.0)
    hdr.columns[0].width = logo_w
    hdr.columns[1].width = mid_w
    hdr.columns[2].width = qr_w

    left, mid, right = hdr.rows[0].cells
    left.width = logo_w
    mid.width = mid_w
    right.width = qr_w

    left.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    mid.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    right.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for c in (left, mid, right):
        _no_cell_border(c)

    # --- Left cell: logo -----------------------------------------------------
    lp = left.paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lp.paragraph_format.space_after = Pt(0)
    logo_file = ASSETS_DIR / "logo.png"
    if logo_file.exists():
        run = lp.add_run()
        run.add_picture(str(logo_file), width=Cm(2.6))

    # --- Middle cell: clinic name / address / phone --------------------------
    mp = mid.paragraphs[0]
    mp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mp.paragraph_format.space_after = Pt(1)
    rn = mp.add_run("{{ clinic.name }}")
    rn.bold = True
    rn.font.size = Pt(13)
    rn.font.color.rgb = RGBColor(0x0A, 0x35, 0x66)

    mp2 = mid.add_paragraph()
    mp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mp2.paragraph_format.space_after = Pt(0)
    ra = mp2.add_run("{{ clinic.address }}")
    ra.font.size = Pt(9)
    ra.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    mp3 = mid.add_paragraph()
    mp3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mp3.paragraph_format.space_after = Pt(0)
    rp = mp3.add_run("Тел: {{ clinic.phone }}")
    rp.font.size = Pt(9)
    rp.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # --- Right cell: QR ------------------------------------------------------
    rrp = right.paragraphs[0]
    rrp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rrp.paragraph_format.space_after = Pt(0)
    qr_file = ASSETS_DIR / "qr.png"
    if qr_file.exists():
        run = rrp.add_run()
        run.add_picture(str(qr_file), width=Cm(2.6))


# ---------------------------------------------------------------------------
# Full document
# ---------------------------------------------------------------------------


def build_default_template(dst: Path | None = None) -> Path:
    """Write the default template to disk and return its path."""
    dst = Path(dst or DEFAULT_DST)

    doc = Document()

    # Tight A4 margins — the whole sheet should fit on one page.
    for section in doc.sections:
        section.top_margin = Cm(1.0)
        section.bottom_margin = Cm(1.0)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)

    # --- Header: two images + clinic identity ------------------------------
    _write_letterhead(doc)
    _add_horizontal_rule(doc, color=BRAND_BLUE_HEX, size=10)

    # --- Title ------------------------------------------------------------
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(2)
    title.paragraph_format.space_after = Pt(1)
    trun = title.add_run("QABUL VARAQASI  /  ЛИСТ ПРИЁМА")
    trun.bold = True
    trun.font.size = Pt(12)
    trun.font.color.rgb = RGBColor(0x0A, 0x35, 0x66)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(4)
    srun = subtitle.add_run("Sana / Дата: {{ reception.date }}    №{{ reception.id }}")
    srun.font.size = Pt(9)
    srun.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # --- Compact patient block --------------------------------------------
    _compact_patient_table(doc)

    # --- Inline sections (compact, no big grey bars) ----------------------
    _inline_section(doc, "SHIKOYATLAR / ЖАЛОБЫ",     "{{ reception.complaints_text }}")
    _inline_section(doc, "ANAMNEZ / АНАМНЕЗ",        "{{ reception.anamnesis }}")
    _inline_section(doc, "LOR STATUS / ЛОР СТАТУС",  "{{ reception.lor_status_text }}")
    _inline_section(doc, "TASHXIS / ДИАГНОЗ",        "{{ reception.diagnosis }}",
                    bold_body=True, body_size=11)

    # Small breather before the framed box
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(2)

    # --- The one framed section — TAVSIYA ---------------------------------
    _boxed_recommendation(doc)

    # --- Signature footer --------------------------------------------------
    doc.add_paragraph()
    sig = doc.add_table(rows=1, cols=2)
    sig.autofit = True
    lc, rc = sig.rows[0].cells
    lc.width = Cm(11.0)
    rc.width = Cm(6.0)

    lp = lc.paragraphs[0]
    lp.paragraph_format.space_after = Pt(0)
    r1 = lp.add_run("Shifokor / Врач: ")
    r1.bold = True
    r1.font.size = Pt(10)
    r2 = lp.add_run("{{ doctor.full_name }}")
    r2.font.size = Pt(10)

    lp2 = lc.add_paragraph()
    lp2.paragraph_format.space_after = Pt(0)
    r3 = lp2.add_run("Tel: {{ doctor.phone }}")
    r3.font.size = Pt(9)
    r3.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    rp = rc.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rp.paragraph_format.space_after = Pt(0)
    r4 = rp.add_run("Imzo: ______________________")
    r4.font.size = Pt(10)

    rp2 = rc.add_paragraph()
    rp2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rp2.paragraph_format.space_after = Pt(0)
    r5 = rp2.add_run("Sana / Дата: {{ today }}")
    r5.font.size = Pt(9)
    r5.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    dst.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(dst))
    return dst


def main() -> None:
    path = build_default_template()
    print(f"[OK] Wrote: {path.relative_to(ROOT)}")


if __name__ == "__main__":
    main()
