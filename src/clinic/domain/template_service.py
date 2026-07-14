"""Management of the shipped Word template file.

The reception .docx renderer looks for a template at
``settings.templates_dir / "reception_template.docx"``. If it exists it is
used as a ``docxtpl`` template; otherwise the renderer falls back to
programmatic ``python-docx`` rendering.

This service exposes a tiny CRUD-ish API around that single file so the
Settings screen can:

- inspect the file's status (exists, size, mtime)
- accept an uploaded replacement
- **append the standard reception sections** to a user's letterhead
  (the operator uploads a clinic-branded letterhead — logo, address,
  phones, QR codes — and the service programmatically adds the patient /
  complaints / diagnosis / signature sections below it)
- reset to the shipped default (regenerates via ``build_reception_template``)
- delete the file entirely (then the fallback renderer kicks in)
"""

from __future__ import annotations

import io
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path

from loguru import logger

from clinic.config import settings
from clinic.printing.docx_builder import TEMPLATE_FILENAME


@dataclass
class TemplateStatus:
    """Snapshot of the currently-installed Word template."""

    exists: bool
    path: Path
    size_bytes: int
    modified: datetime | None

    @property
    def size_kb(self) -> float:
        return round(self.size_bytes / 1024.0, 1) if self.size_bytes else 0.0


# Hard cap on uploaded template size — keeps upload handling simple. A
# realistic reception template is < 100 KB; 5 MB is a very generous limit
# that still guards against accidental huge uploads.
MAX_TEMPLATE_BYTES = 5 * 1024 * 1024

# docx MIME + magic bytes. .docx is a ZIP archive so the first two bytes
# are "PK". We use the magic check because Windows browsers often send
# ``application/octet-stream`` or a Word-specific mimetype rather than the
# ``vnd.openxmlformats`` MIME.
_DOCX_MAGIC = b"PK\x03\x04"


class TemplateError(Exception):
    """Raised for validation failures on template upload."""


def template_path() -> Path:
    return Path(settings.templates_dir) / TEMPLATE_FILENAME


def status() -> TemplateStatus:
    p = template_path()
    if not p.is_file():
        return TemplateStatus(exists=False, path=p, size_bytes=0, modified=None)
    st = p.stat()
    return TemplateStatus(
        exists=True,
        path=p,
        size_bytes=st.st_size,
        modified=datetime.fromtimestamp(st.st_mtime),
    )


def save_uploaded(payload: bytes) -> TemplateStatus:
    """Validate + persist an uploaded .docx as the current template.

    Raises :class:`TemplateError` on any validation issue.
    """
    if not payload:
        raise TemplateError("empty")
    if len(payload) > MAX_TEMPLATE_BYTES:
        raise TemplateError("too_large")
    if not payload.startswith(_DOCX_MAGIC):
        raise TemplateError("invalid")

    # Round-trip through python-docx to make sure the file is a readable
    # Word document. A raw ZIP will pass the magic check but might not be
    # a valid .docx.
    import io

    from docx import Document as _Doc

    try:
        _ = _Doc(io.BytesIO(payload))
    except Exception as exc:
        logger.warning("Rejected uploaded template: {}", exc)
        raise TemplateError("invalid") from exc

    p = template_path()
    p.parent.mkdir(parents=True, exist_ok=True)
    p.write_bytes(payload)
    logger.info("Saved uploaded template ({} bytes) to {}", len(payload), p)
    return status()


def reset_to_default() -> TemplateStatus:
    """Rewrite the template file from the shipped baseline."""
    from scripts.build_reception_template import build_default_template

    build_default_template(template_path())
    logger.info("Reception template reset to shipped default")
    return status()


def delete() -> bool:
    """Remove the template file. Returns True if a file was removed."""
    p = template_path()
    if not p.is_file():
        return False
    p.unlink()
    logger.info("Reception template deleted")
    return True


# ---------------------------------------------------------------------------
# Letterhead extension
# ---------------------------------------------------------------------------


def save_letterhead_with_sections(payload: bytes) -> TemplateStatus:
    """Accept a letterhead-only docx and save it as a full template.

    The operator uploads *only* their clinic-branded letterhead (logo, name,
    address, phones, QR code — the top-of-page block). This function opens
    that file, appends the standard reception sections (patient block +
    complaints + anamnesis + LOR STATUS + diagnosis + recommendations +
    signature) with the appropriate ``docxtpl`` placeholders, and saves the
    result as the active template.

    Raises :class:`TemplateError` on any validation issue.
    """
    if not payload:
        raise TemplateError("empty")
    if len(payload) > MAX_TEMPLATE_BYTES:
        raise TemplateError("too_large")
    if not payload.startswith(_DOCX_MAGIC):
        raise TemplateError("invalid")

    from docx import Document as _Doc

    try:
        doc = _Doc(io.BytesIO(payload))
    except Exception as exc:
        logger.warning("Rejected uploaded letterhead: {}", exc)
        raise TemplateError("invalid") from exc

    # Append the reception sections in-place.
    _append_reception_sections(doc)

    # Serialise back to bytes and persist.
    buf = io.BytesIO()
    doc.save(buf)
    out = buf.getvalue()

    p = template_path()
    p.parent.mkdir(parents=True, exist_ok=True)
    p.write_bytes(out)
    logger.info(
        "Letterhead extended with reception sections ({} bytes in, {} bytes out) -> {}",
        len(payload),
        len(out),
        p,
    )
    return status()


def _append_reception_sections(doc) -> None:
    """Add the standard reception sections at the end of ``doc``.

    The insertion always happens *after* the existing content, so the
    operator's letterhead (header/footer/logos/QR/branding text) is
    preserved verbatim. Everything below uses docxtpl placeholders so the
    reception renderer can fill them in at print time.
    """
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor

    # --- helpers ----------------------------------------------------------

    def _shade(cell, hex_color: str) -> None:
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color)
        tc_pr.append(shd)

    def _thin_border(cell, size: int = 4, color: str = "808080") -> None:
        tc_pr = cell._tc.get_or_add_tcPr()
        borders = tc_pr.find(qn("w:tcBorders"))
        if borders is None:
            borders = OxmlElement("w:tcBorders")
            tc_pr.append(borders)
        for edge in ("top", "left", "bottom", "right"):
            existing = borders.find(qn(f"w:{edge}"))
            if existing is not None:
                borders.remove(existing)
            el = OxmlElement(f"w:{edge}")
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), str(size))
            el.set(qn("w:color"), color)
            borders.append(el)

    def _add_divider(doc_) -> None:
        p = doc_.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(2)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "8")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "0A3566")
        pBdr.append(bottom)
        pPr.append(pBdr)

    def _add_heading(doc_, text: str) -> None:
        p = doc_.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x11, 0x11, 0x11)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "808080")
        pBdr.append(bottom)
        pPr.append(pBdr)

    def _add_body(doc_, placeholder: str, *, bold: bool = False, size: int = 11) -> None:
        p = doc_.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(placeholder)
        run.font.size = Pt(size)
        run.bold = bold

    # --- start layout below the letterhead --------------------------------

    _add_divider(doc)

    # Document title.
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(4)
    title.paragraph_format.space_after = Pt(2)
    trun = title.add_run("QABUL VARAQASI  /  ЛИСТ ПРИЁМА")
    trun.bold = True
    trun.font.size = Pt(13)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(10)
    srun = subtitle.add_run("Sana / Дата: {{ reception.date }}    №{{ reception.id }}")
    srun.font.size = Pt(10)
    srun.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # Patient block — 2-column key/value table.
    _add_heading(doc, "BEMOR MA'LUMOTLARI  /  СВЕДЕНИЯ О ПАЦИЕНТЕ")

    rows = [
        ("F.I.SH. / Ф.И.О.", "{{ patient.full_name }}"),
        ("Tug'ilgan yili / Год рождения",
         "{{ patient.birth_year }}  ({{ patient.age }} yosh / лет)"),
        ("Manzil / Адрес", "{{ patient.address }}"),
        ("Telefon / Телефон", "{{ patient.phone }}"),
    ]
    tbl = doc.add_table(rows=len(rows), cols=2)
    tbl.autofit = False
    for i, (k, v) in enumerate(rows):
        c0, c1 = tbl.rows[i].cells
        c0.width = Cm(5.0)
        c1.width = Cm(12.0)
        c0.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        c1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _shade(c0, "F5F5F5")
        r0 = c0.paragraphs[0].add_run(k)
        r0.bold = True
        r0.font.size = Pt(10)
        r0.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
        r1 = c1.paragraphs[0].add_run(v)
        r1.font.size = Pt(11)
        _thin_border(c0)
        _thin_border(c1)

    # Content sections.
    _add_heading(doc, "SHIKOYATLAR  /  ЖАЛОБЫ")
    _add_body(doc, "{{ reception.complaints_text }}")

    _add_heading(doc, "ANAMNEZ  /  АНАМНЕЗ")
    _add_body(doc, "{{ reception.anamnesis }}")

    _add_heading(doc, "LOR STATUS  /  ЛОР СТАТУС")
    _add_body(doc, "{{ reception.lor_status_text }}")

    _add_heading(doc, "TASHXIS  /  ДИАГНОЗ")
    _add_body(doc, "{{ reception.diagnosis }}", bold=True, size=12)

    _add_heading(doc, "TAVSIYALAR  /  РЕКОМЕНДАЦИИ")
    _add_body(doc, "{{ reception.recommendation }}")

    # Doctor signature block — 2-column table, no borders.
    doc.add_paragraph()  # spacer
    sig_tbl = doc.add_table(rows=1, cols=2)
    sig_tbl.autofit = True
    left, right = sig_tbl.rows[0].cells

    lp = left.paragraphs[0]
    r1 = lp.add_run("Shifokor / Врач: ")
    r1.bold = True
    r1.font.size = Pt(11)
    r2 = lp.add_run("{{ doctor.full_name }}")
    r2.font.size = Pt(11)

    lp2 = left.add_paragraph()
    r3 = lp2.add_run("Tel: {{ doctor.phone }}")
    r3.font.size = Pt(10)
    r3.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    rp = right.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r4 = rp.add_run("Imzo / Подпись: ______________________")
    r4.font.size = Pt(11)

    rp2 = right.add_paragraph()
    rp2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r5 = rp2.add_run("Sana / Дата: {{ today }}")
    r5.font.size = Pt(10)
    r5.font.color.rgb = RGBColor(0x60, 0x60, 0x60)


__all__ = [
    "MAX_TEMPLATE_BYTES",
    "TemplateError",
    "TemplateStatus",
    "delete",
    "reset_to_default",
    "save_letterhead_with_sections",
    "save_uploaded",
    "status",
    "template_path",
]
