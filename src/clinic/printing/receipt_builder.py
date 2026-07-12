"""Build a Word ``.docx`` receipt for a cashier payment.

Layout:

* Clinic identity header (name + address + phone).
* Title (``KVITANSIYA`` / ``КВИТАНЦИЯ``).
* Patient block + payment metadata (date, receipt number).
* Table of services (name, quantity, unit price, line total).
* Grand total.
* Optional user note.
* Signature line.

The design mirrors :mod:`docx_builder` — supports a user template via
``templates/receipt_template.docx`` and falls back to a from-scratch renderer
when no template is present.
"""

from __future__ import annotations

from decimal import Decimal
from pathlib import Path
from typing import Any

from docx import Document
from docx.document import Document as DocumentType
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor
from docxtpl import DocxTemplate
from loguru import logger

from clinic.config import settings
from clinic.domain.dto import CashierRecordDTO, PatientDTO
from clinic.printing.common import format_date, format_money

TEMPLATE_FILENAME = "receipt_template.docx"


# ============================================================================
# Public API
# ============================================================================


def build_receipt_document(
    *,
    records: list[CashierRecordDTO],
    patient: PatientDTO,
    clinic: dict[str, str],
    lang: str = "uz",
    template_path: Path | None = None,
) -> DocumentType:
    if not records:
        raise ValueError("Cannot build a receipt without records")

    context = _build_context(records=records, patient=patient, clinic=clinic, lang=lang)
    template = _resolve_template(template_path)
    if template is not None:
        return _render_with_template(template, context)
    return _render_default(context, lang=lang)


def save_receipt_document(output_path: Path, **kwargs: Any) -> Path:
    doc = build_receipt_document(**kwargs)
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    logger.info("Receipt saved to: {}", output_path)
    return output_path


# ============================================================================
# Context
# ============================================================================


_LABELS = {
    "title": {"uz": "KVITANSIYA", "ru": "КВИТАНЦИЯ"},
    "patient": {"uz": "Bemor", "ru": "Пациент"},
    "date": {"uz": "Sana", "ru": "Дата"},
    "receipt_no": {"uz": "Kvitansiya #", "ru": "Квитанция №"},
    "col_num": {"uz": "#", "ru": "#"},
    "col_service": {"uz": "Xizmat", "ru": "Услуга"},
    "col_qty": {"uz": "Soni", "ru": "Кол-во"},
    "col_price": {"uz": "Narx", "ru": "Цена"},
    "col_total": {"uz": "Jami", "ru": "Сумма"},
    "grand_total": {"uz": "JAMI", "ru": "ИТОГО"},
    "note": {"uz": "Izoh", "ru": "Примечание"},
    "signature": {"uz": "Kassir imzosi", "ru": "Подпись кассира"},
    "currency": {"uz": "so'm", "ru": "сум"},
}


def _label(key: str, lang: str) -> str:
    entry = _LABELS.get(key, {})
    return entry.get(lang) or entry.get("uz") or key


def _build_context(
    *,
    records: list[CashierRecordDTO],
    patient: PatientDTO,
    clinic: dict[str, str],
    lang: str,
) -> dict[str, Any]:
    items = []
    grand = Decimal("0")
    note_parts: list[str] = []
    for i, r in enumerate(records, start=1):
        items.append(
            {
                "num": i,
                "service": r.service_name(lang),
                "quantity": r.quantity,
                "price": format_money(r.price_at_moment),
                "total": format_money(r.total),
            }
        )
        grand += r.total
        if r.note and r.note.strip() and r.note not in note_parts:
            note_parts.append(r.note.strip())

    receipt_no = records[0].id if records else 0
    paid_at = records[0].paid_at

    return {
        "clinic": {
            "name": clinic.get(f"name_{lang}") or clinic.get("name_uz") or "",
            "address": clinic.get(f"address_{lang}") or clinic.get("address_uz") or "",
            "phone": clinic.get("phone", ""),
            "logo_path": clinic.get("logo_path", ""),
        },
        "patient": {
            "id": patient.id,
            "full_name": patient.full_name,
            "birth_year": patient.birth_year,
            "phone": patient.phone or "",
        },
        "receipt": {
            "id": receipt_no,
            "date": format_date(paid_at, with_time=True),
            "note": " · ".join(note_parts) if note_parts else "",
        },
        "items": items,
        "grand_total": format_money(grand),
        "currency": _label("currency", lang),
        "lang": lang,
    }


# ============================================================================
# Template rendering
# ============================================================================


def _resolve_template(explicit: Path | None) -> Path | None:
    if explicit:
        p = Path(explicit)
        return p if p.is_file() else None
    candidate = settings.templates_dir / TEMPLATE_FILENAME
    return candidate if candidate.is_file() else None


def _render_with_template(template_path: Path, context: dict[str, Any]) -> DocumentType:
    tpl = DocxTemplate(str(template_path))
    tpl.render(context)
    import io

    buf = io.BytesIO()
    tpl.save(buf)
    buf.seek(0)
    return Document(buf)


# ============================================================================
# Default rendering
# ============================================================================


def _render_default(context: dict[str, Any], *, lang: str) -> DocumentType:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)

    clinic = context.get("clinic", {})
    patient = context.get("patient", {})
    receipt = context.get("receipt", {})
    currency = context.get("currency", "")

    # Header
    if clinic.get("name"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(clinic["name"])
        run.bold = True
        run.font.size = Pt(14)

    subline = " · ".join(x for x in (clinic.get("address"), clinic.get("phone")) if x)
    if subline:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(subline)
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    doc.add_paragraph()

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(_label("title", lang))
    run.bold = True
    run.font.size = Pt(13)

    doc.add_paragraph()

    # Meta lines
    _kv_line(doc, _label("patient", lang), patient.get("full_name", "—"))
    _kv_line(doc, _label("date", lang), receipt.get("date", "—"))
    _kv_line(doc, _label("receipt_no", lang), str(receipt.get("id", "—")))

    doc.add_paragraph()

    # Items table
    items = context.get("items", [])
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, key in enumerate(
        ("col_num", "col_service", "col_qty", "col_price", "col_total")
    ):
        cell = hdr[i]
        p = cell.paragraphs[0]
        run = p.add_run(_label(key, lang))
        run.bold = True
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for row in items:
        r = table.add_row().cells
        r[0].text = str(row["num"])
        r[1].text = row["service"]
        r[2].text = str(row["quantity"])
        r[3].text = f"{row['price']} {currency}"
        r[4].text = f"{row['total']} {currency}"
        r[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        r[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        r[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.add_paragraph()

    # Grand total
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(f"{_label('grand_total', lang)}: {context.get('grand_total', '0')} {currency}")
    run.bold = True
    run.font.size = Pt(13)

    # Note (optional)
    if receipt.get("note"):
        _kv_line(doc, _label("note", lang), receipt["note"])

    doc.add_paragraph()
    doc.add_paragraph()

    # Signature
    p = doc.add_paragraph()
    p.add_run(f"{_label('signature', lang)}: ______________________")

    return doc


def _kv_line(doc: DocumentType, label: str, value: str) -> None:
    p = doc.add_paragraph()
    label_run = p.add_run(f"{label}: ")
    label_run.bold = True
    p.add_run(value if value else "—")


__all__ = ["TEMPLATE_FILENAME", "build_receipt_document", "save_receipt_document"]
