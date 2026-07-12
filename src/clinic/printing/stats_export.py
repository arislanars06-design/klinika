"""Word export for patient + cashier statistics dashboards."""

from __future__ import annotations

from decimal import Decimal
from pathlib import Path
from typing import Any

from docx import Document
from docx.document import Document as DocumentType
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor
from docxtpl import DocxTemplate
from loguru import logger

from clinic.config import settings
from clinic.domain.dto import CashierStats, PatientStats
from clinic.domain.stats_service import Period
from clinic.printing.common import format_date, format_money

PATIENTS_TEMPLATE_FILENAME = "patients_stats_template.docx"
CASHIER_TEMPLATE_FILENAME = "cashier_stats_template.docx"


# ============================================================================
# Public API
# ============================================================================


def build_patient_stats_document(
    stats: PatientStats,
    period: Period,
    *,
    clinic: dict[str, str],
    lang: str = "uz",
    template_path: Path | None = None,
) -> DocumentType:
    context = _patient_context(stats, period, clinic, lang)
    template = _resolve(template_path, PATIENTS_TEMPLATE_FILENAME)
    if template is not None:
        return _render_with_template(template, context)
    return _render_patient_default(context, lang=lang)


def build_cashier_stats_document(
    stats: CashierStats,
    period: Period,
    *,
    clinic: dict[str, str],
    lang: str = "uz",
    template_path: Path | None = None,
) -> DocumentType:
    context = _cashier_context(stats, period, clinic, lang)
    template = _resolve(template_path, CASHIER_TEMPLATE_FILENAME)
    if template is not None:
        return _render_with_template(template, context)
    return _render_cashier_default(context, lang=lang)


def save_patient_stats(output_path: Path, **kwargs: Any) -> Path:
    doc = build_patient_stats_document(**kwargs)
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    logger.info("Patient stats saved to: {}", output_path)
    return output_path


def save_cashier_stats(output_path: Path, **kwargs: Any) -> Path:
    doc = build_cashier_stats_document(**kwargs)
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    logger.info("Cashier stats saved to: {}", output_path)
    return output_path


# ============================================================================
# Labels + contexts
# ============================================================================


_LABELS = {
    "period": {"uz": "Davr", "ru": "Период"},
    "patients_title": {"uz": "KLINIKA STATISTIKASI", "ru": "СТАТИСТИКА КЛИНИКИ"},
    "cashier_title": {"uz": "KASSA HISOBOTI", "ru": "КАССОВЫЙ ОТЧЁТ"},
    "kpi_total": {"uz": "Jami bemorlar (noyob)", "ru": "Всего пациентов (уник.)"},
    "kpi_new": {"uz": "Yangi bemorlar", "ru": "Новых пациентов"},
    "kpi_repeat": {"uz": "Takroriy qabullar", "ru": "Повторных приёмов"},
    "kpi_revenue": {"uz": "Jami tushum", "ru": "Общая выручка"},
    "kpi_receipts": {"uz": "To'lovlar (chek) soni", "ru": "Оплат (чеков)"},
    "kpi_avg": {"uz": "O'rtacha chek", "ru": "Средний чек"},
    "top_diagnoses": {"uz": "TOP tashxislar", "ru": "TOP диагнозов"},
    "by_service": {"uz": "Xizmatlar bo'yicha", "ru": "По услугам"},
    "by_day": {"uz": "Kunlar bo'yicha", "ru": "По дням"},
    "col_diagnosis": {"uz": "Tashxis", "ru": "Диагноз"},
    "col_count": {"uz": "Soni", "ru": "Кол-во"},
    "col_service": {"uz": "Xizmat", "ru": "Услуга"},
    "col_units": {"uz": "Soni", "ru": "Кол-во"},
    "col_revenue": {"uz": "Tushum", "ru": "Выручка"},
    "col_date": {"uz": "Sana", "ru": "Дата"},
    "col_value": {"uz": "Qiymat", "ru": "Значение"},
    "no_data": {"uz": "Ma'lumot yo'q", "ru": "Нет данных"},
    "currency": {"uz": "so'm", "ru": "сум"},
}


def _label(key: str, lang: str) -> str:
    entry = _LABELS.get(key, {})
    return entry.get(lang) or entry.get("uz") or key


def _period_label(period: Period) -> str:
    return f"{format_date(period.start)} — {format_date(period.end)}"


def _clinic_context(clinic: dict[str, str], lang: str) -> dict[str, str]:
    return {
        "name": clinic.get(f"name_{lang}") or clinic.get("name_uz") or "",
        "address": clinic.get(f"address_{lang}") or clinic.get("address_uz") or "",
        "phone": clinic.get("phone", ""),
        "logo_path": clinic.get("logo_path", ""),
    }


def _patient_context(
    stats: PatientStats,
    period: Period,
    clinic: dict[str, str],
    lang: str,
) -> dict[str, Any]:
    return {
        "clinic": _clinic_context(clinic, lang),
        "period": {
            "start": format_date(period.start),
            "end": format_date(period.end),
            "label": _period_label(period),
        },
        "kpis": {
            "total_patients": stats.total_patients,
            "new_patients": stats.new_patients,
            "repeat_receptions": stats.repeat_receptions,
        },
        "top_diagnoses": [
            {"diagnosis": d.diagnosis, "count": d.count} for d in stats.top_diagnoses
        ],
        "by_day": [{"date": p.date, "value": p.value} for p in stats.by_day],
        "lang": lang,
    }


def _cashier_context(
    stats: CashierStats,
    period: Period,
    clinic: dict[str, str],
    lang: str,
) -> dict[str, Any]:
    return {
        "clinic": _clinic_context(clinic, lang),
        "period": {
            "start": format_date(period.start),
            "end": format_date(period.end),
            "label": _period_label(period),
        },
        "kpis": {
            "total_revenue": format_money(stats.total_revenue),
            "payment_count": stats.payment_count,
            "receipts_count": stats.receipts_count,
            "average_receipt": format_money(stats.average_receipt),
        },
        "by_service": [
            {
                "service": svc.display_name(lang),
                "units": svc.units_sold,
                "revenue": format_money(svc.revenue),
            }
            for svc in stats.by_service
        ],
        "by_day": [
            {"date": p.date, "value": format_money(Decimal(str(p.value)))}
            for p in stats.by_day
        ],
        "currency": _label("currency", lang),
        "lang": lang,
    }


# ============================================================================
# Template rendering
# ============================================================================


def _resolve(explicit: Path | None, default_name: str) -> Path | None:
    if explicit:
        p = Path(explicit)
        return p if p.is_file() else None
    candidate = settings.templates_dir / default_name
    return candidate if candidate.is_file() else None


def _render_with_template(template_path: Path, context: dict[str, Any]) -> DocumentType:
    tpl = DocxTemplate(str(template_path))
    tpl.render(context)
    import io

    buf = io.BytesIO()
    tpl.save(buf)
    buf.seek(0)
    return Document(buf)


# ============================================================================
# Default (from-scratch) rendering
# ============================================================================


def _render_header(doc: DocumentType, context: dict[str, Any], title: str) -> None:
    clinic = context.get("clinic", {})
    for section in doc.sections:
        section.top_margin = Cm(1.6)
        section.bottom_margin = Cm(1.6)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)

    if clinic.get("name"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(clinic["name"])
        run.bold = True
        run.font.size = Pt(14)

    subline = " · ".join(x for x in (clinic.get("address"), clinic.get("phone")) if x)
    if subline:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(subline)
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(14)

    period = context.get("period", {})
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"{context.get('period_label_prefix', '')}{period.get('label', '')}").italic = True

    doc.add_paragraph()


def _render_kpi_block(doc: DocumentType, kpis: list[tuple[str, str]]) -> None:
    if not kpis:
        return
    table = doc.add_table(rows=1, cols=len(kpis))
    table.style = "Light Shading Accent 1"
    for i, (label, value) in enumerate(kpis):
        cell = table.rows[0].cells[i]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(str(value))
        run.bold = True
        run.font.size = Pt(14)
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = p2.add_run(label)
        run2.font.size = Pt(9)
        run2.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
    doc.add_paragraph()


def _render_two_col_table(
    doc: DocumentType,
    *,
    title: str,
    headers: tuple[str, str],
    rows: list[tuple[str, str]],
    empty_label: str,
) -> None:
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(12)

    if not rows:
        doc.add_paragraph(empty_label)
        return

    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.paragraphs[0].add_run(header).bold = True

    for a, b in rows:
        r = table.add_row().cells
        r[0].text = str(a)
        r[1].text = str(b)
        r[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_paragraph()


def _render_three_col_table(
    doc: DocumentType,
    *,
    title: str,
    headers: tuple[str, str, str],
    rows: list[tuple[str, str, str]],
    empty_label: str,
) -> None:
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(12)

    if not rows:
        doc.add_paragraph(empty_label)
        return

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.paragraphs[0].add_run(header).bold = True

    for a, b, c in rows:
        r = table.add_row().cells
        r[0].text = str(a)
        r[1].text = str(b)
        r[2].text = str(c)
        r[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_paragraph()


def _render_patient_default(context: dict[str, Any], *, lang: str) -> DocumentType:
    doc = Document()
    _render_header(doc, context, _label("patients_title", lang))
    empty = _label("no_data", lang)

    kpis = context.get("kpis", {})
    _render_kpi_block(
        doc,
        [
            (_label("kpi_total", lang), kpis.get("total_patients", 0)),
            (_label("kpi_new", lang), kpis.get("new_patients", 0)),
            (_label("kpi_repeat", lang), kpis.get("repeat_receptions", 0)),
        ],
    )

    _render_two_col_table(
        doc,
        title=_label("top_diagnoses", lang),
        headers=(_label("col_diagnosis", lang), _label("col_count", lang)),
        rows=[(row["diagnosis"], str(row["count"])) for row in context.get("top_diagnoses", [])],
        empty_label=empty,
    )

    _render_two_col_table(
        doc,
        title=_label("by_day", lang),
        headers=(_label("col_date", lang), _label("col_count", lang)),
        rows=[(row["date"], str(int(row["value"]))) for row in context.get("by_day", [])],
        empty_label=empty,
    )
    return doc


def _render_cashier_default(context: dict[str, Any], *, lang: str) -> DocumentType:
    doc = Document()
    _render_header(doc, context, _label("cashier_title", lang))
    empty = _label("no_data", lang)
    currency = context.get("currency", "")

    kpis = context.get("kpis", {})
    _render_kpi_block(
        doc,
        [
            (_label("kpi_revenue", lang), f"{kpis.get('total_revenue', 0)} {currency}"),
            (_label("kpi_receipts", lang), kpis.get("receipts_count", 0)),
            (_label("kpi_avg", lang), f"{kpis.get('average_receipt', 0)} {currency}"),
        ],
    )

    _render_three_col_table(
        doc,
        title=_label("by_service", lang),
        headers=(
            _label("col_service", lang),
            _label("col_units", lang),
            _label("col_revenue", lang),
        ),
        rows=[
            (row["service"], str(row["units"]), f"{row['revenue']} {currency}")
            for row in context.get("by_service", [])
        ],
        empty_label=empty,
    )

    _render_two_col_table(
        doc,
        title=_label("by_day", lang),
        headers=(_label("col_date", lang), _label("col_revenue", lang)),
        rows=[
            (row["date"], f"{row['value']} {currency}")
            for row in context.get("by_day", [])
        ],
        empty_label=empty,
    )
    return doc


__all__ = [
    "CASHIER_TEMPLATE_FILENAME",
    "PATIENTS_TEMPLATE_FILENAME",
    "build_cashier_stats_document",
    "build_patient_stats_document",
    "save_cashier_stats",
    "save_patient_stats",
]
