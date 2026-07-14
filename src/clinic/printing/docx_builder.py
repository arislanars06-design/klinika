"""Build a Word ``.docx`` for one reception.

Two rendering strategies are supported:

1. **User template** — if ``templates/reception_template.docx`` exists (or an
   explicit path is provided), it is treated as a docxtpl / Jinja template.
2. **Default rendering** — otherwise, we build the document from scratch with
   ``python-docx`` so the app still produces a printable page out of the box.

Both paths share the same context dictionary so a user template stays
compatible without touching Python code.

Reference for the placeholder catalog: ``docs/template_placeholders.md``.
"""

from __future__ import annotations

from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.document import Document as DocumentType
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Mm, Pt, RGBColor
from docxtpl import DocxTemplate, InlineImage
from loguru import logger

from clinic.config import settings
from clinic.domain.dto import DoctorDTO, PatientDTO, ReceptionDTO
from clinic.printing.common import format_age, format_date
from clinic.printing.text_composer import compose_complaints, compose_lor_status

TEMPLATE_FILENAME = "reception_template.docx"


# ============================================================================
# Public API
# ============================================================================


def build_reception_document(
    *,
    reception: ReceptionDTO,
    patient: PatientDTO,
    doctor: DoctorDTO | None,
    clinic: dict[str, str],
    lang: str = "uz",
    template_path: Path | None = None,
) -> DocumentType:
    """Return a ``docx.Document`` populated with reception data."""
    context = _build_context(
        reception=reception,
        patient=patient,
        doctor=doctor,
        clinic=clinic,
        lang=lang,
    )
    template = _resolve_template(template_path)
    if template is not None:
        return _render_with_template(template, context)
    return _render_default(context, lang=lang)


def save_reception_document(
    output_path: Path,
    **kwargs: Any,
) -> Path:
    """Convenience wrapper: build + save to disk. Returns the written path."""
    doc = build_reception_document(**kwargs)
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    logger.info("Reception document saved to: {}", output_path)
    return output_path


# ============================================================================
# Context building
# ============================================================================


def _build_context(
    *,
    reception: ReceptionDTO,
    patient: PatientDTO,
    doctor: DoctorDTO | None,
    clinic: dict[str, str],
    lang: str,
) -> dict[str, Any]:
    complaints_text = compose_complaints(
        reception.complaints_codes,
        reception.complaints_details,
        reception.complaints_note,
        lang=lang,
    )
    lor_status_text = compose_lor_status(reception.lor_status, lang=lang)

    return {
        "clinic": {
            "name": clinic.get(f"name_{lang}")
            or clinic.get("name_uz")
            or "",
            "address": clinic.get(f"address_{lang}")
            or clinic.get("address_uz")
            or "",
            "phone": clinic.get("phone") or "",
            "logo_path": clinic.get("logo_path") or "",
        },
        "patient": {
            "id": patient.id,
            "full_name": patient.full_name,
            "birth_year": patient.birth_year,
            "age": format_age(patient.birth_year),
            "address": patient.address or "",
            "phone": patient.phone or "",
        },
        "reception": {
            "id": reception.id,
            "date": format_date(reception.reception_date, with_time=True),
            "complaints_text": complaints_text,
            "complaints_codes": list(reception.complaints_codes or []),
            "complaints_note": reception.complaints_note or "",
            "anamnesis": reception.anamnesis or "",
            "lor_status_text": lor_status_text,
            "diagnosis": reception.diagnosis,
            "recommendation": reception.recommendation or "",
        },
        "doctor": {
            "full_name": doctor.full_name if doctor else "",
            "phone": (doctor.phone if doctor and doctor.phone else ""),
        },
        "today": format_date(datetime.now(), with_time=False),
        "lang": lang,
    }


# ============================================================================
# Template path resolution
# ============================================================================


def _resolve_template(explicit: Path | None) -> Path | None:
    """Return an existing template path or ``None``."""
    if explicit:
        p = Path(explicit)
        return p if p.is_file() else None
    candidate = settings.templates_dir / TEMPLATE_FILENAME
    return candidate if candidate.is_file() else None


def _render_with_template(template_path: Path, context: dict[str, Any]) -> DocumentType:
    tpl = DocxTemplate(str(template_path))
    # Optional inline logo — templates may include ``{{ clinic.logo }}``.
    logo_path = context.get("clinic", {}).get("logo_path")
    if logo_path and Path(logo_path).is_file():
        try:
            context["clinic"]["logo"] = InlineImage(
                tpl, str(logo_path), width=Mm(25)
            )
        except Exception:
            logger.exception("Failed to embed logo in template")
    tpl.render(context)
    # DocxTemplate does not directly expose a Document instance to callers.
    # We save to an in-memory buffer and reload so downstream code gets a
    # regular ``Document``.
    import io

    buf = io.BytesIO()
    tpl.save(buf)
    buf.seek(0)
    return Document(buf)


# ============================================================================
# Default (template-less) rendering
# ============================================================================


_HEADING_UZ = "ҚАБУЛ ВАРАҚАСИ"
_HEADING_RU = "ЛИСТ ПРИЁМА"

# Uzbek labels are Cyrillic to match the rest of the UI. LOR STATUS is
# deliberately kept in Latin — the LOR terminology is Latin-only per user
# request.
_SECTION_LABELS = {
    "patient": {"uz": "БЕМОР", "ru": "ПАЦИЕНТ"},
    "reception_date": {"uz": "Қабул санаси", "ru": "Дата приёма"},
    "birth_year": {"uz": "Туғилган йили", "ru": "Год рождения"},
    "age_suffix": {"uz": "ёш", "ru": "лет"},
    "address": {"uz": "Манзил", "ru": "Адрес"},
    "phone": {"uz": "Телефон", "ru": "Телефон"},
    "complaints": {"uz": "ШИКОЯТЛАР", "ru": "ЖАЛОБЫ"},
    "anamnesis": {"uz": "АНАМНЕЗ", "ru": "АНАМНЕЗ"},
    "lor_status": {"uz": "LOR STATUS", "ru": "LOR STATUS"},
    "diagnosis": {"uz": "ТАШХИС", "ru": "ДИАГНОЗ"},
    "recommendation": {"uz": "ТАВСИЯ", "ru": "РЕКОМЕНДАЦИИ"},
    "doctor": {"uz": "Шифокор", "ru": "Врач"},
    "signature": {"uz": "Имзо", "ru": "Подпись"},
    "date_short": {"uz": "Сана", "ru": "Дата"},
}


def _label(key: str, lang: str) -> str:
    entry = _SECTION_LABELS.get(key, {})
    return entry.get(lang) or entry.get("uz") or key


def _render_default(context: dict[str, Any], *, lang: str) -> DocumentType:
    doc = Document()

    # Page margins tuned for A4 portrait
    for section in doc.sections:
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    clinic = context.get("clinic", {})
    patient = context.get("patient", {})
    reception = context.get("reception", {})
    doctor = context.get("doctor", {})

    # ----- header (clinic identity) -----
    if clinic.get("name"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(clinic["name"])
        run.bold = True
        run.font.size = Pt(16)
    subline_parts = [clinic.get("address", ""), clinic.get("phone", "")]
    subline = " · ".join(x for x in subline_parts if x)
    if subline:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(subline)
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    doc.add_paragraph()

    # ----- title -----
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(_HEADING_RU if lang == "ru" else _HEADING_UZ)
    run.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph()

    # ----- patient block: 4-row two-column table -----
    patient_rows = [
        (_label("patient", lang), patient.get("full_name", "—")),
        (
            _label("birth_year", lang),
            f"{patient.get('birth_year', '—')} ({patient.get('age', '')} {_label('age_suffix', lang)})",
        ),
        (_label("address", lang), patient.get("address") or "—"),
        (
            f"{_label('phone', lang)} / {_label('reception_date', lang)}",
            f"{patient.get('phone') or '—'}  ·  {reception.get('date', '—')}",
        ),
    ]
    table = doc.add_table(rows=len(patient_rows), cols=2)
    table.autofit = False
    for row_idx, (label, value) in enumerate(patient_rows):
        cells = table.rows[row_idx].cells
        cells[0].width = Cm(5.0)
        cells[1].width = Cm(11.0)
        _cell_text(cells[0], label, bold=True)
        _cell_text(cells[1], value, bold=(row_idx == 0))

    # ----- complaints -----
    _add_heading(doc, _label("complaints", lang))
    _add_paragraph(doc, reception.get("complaints_text") or "—")

    # ----- anamnesis -----
    if reception.get("anamnesis"):
        _add_heading(doc, _label("anamnesis", lang))
        _add_paragraph(doc, reception["anamnesis"])

    # ----- LOR STATUS -----
    if reception.get("lor_status_text"):
        _add_heading(doc, _label("lor_status", lang))
        _add_paragraph(doc, reception["lor_status_text"])

    # ----- diagnosis -----
    _add_heading(doc, _label("diagnosis", lang))
    p = doc.add_paragraph()
    run = p.add_run(reception.get("diagnosis") or "—")
    run.bold = True
    run.font.size = Pt(12)

    # ----- recommendation -----
    if reception.get("recommendation"):
        _add_heading(doc, _label("recommendation", lang))
        _add_paragraph(doc, reception["recommendation"])

    doc.add_paragraph()
    doc.add_paragraph()

    # ----- doctor signature line -----
    p = doc.add_paragraph()
    left = p.add_run(f"{_label('doctor', lang)}: {doctor.get('full_name', '—')}")
    left.font.size = Pt(11)
    p.add_run("\t" * 4)
    p.add_run(f"{_label('signature', lang)}: ______________________")

    p = doc.add_paragraph()
    right = p.add_run(f"{_label('date_short', lang)}: {context.get('today', '')}")
    right.font.size = Pt(10)
    right.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    return doc


# ============================================================================
# python-docx helpers
# ============================================================================


def _add_kv_line(doc: DocumentType, label: str, value: str, *, bold: bool = False) -> None:
    p = doc.add_paragraph()
    label_run = p.add_run(f"{label}: ")
    label_run.bold = True
    value_run = p.add_run(value if value else "—")
    if bold:
        value_run.bold = True


def _add_heading(doc: DocumentType, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x21, 0x21, 0x21)


def _add_paragraph(doc: DocumentType, text: str) -> None:
    for line in text.split("\n"):
        doc.add_paragraph(line)


def _cell_text(cell, text: str, *, bold: bool = False) -> None:
    """Write ``text`` into a table cell with a specific run style."""
    # ``add_table`` seeds each cell with an empty paragraph — reuse it so we
    # don't end up with a leading blank line.
    para = cell.paragraphs[0]
    run = para.add_run(text if text else "—")
    run.bold = bold
    run.font.size = Pt(11)


__all__ = [
    "TEMPLATE_FILENAME",
    "build_reception_document",
    "save_reception_document",
]
